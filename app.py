"""
Month-End Close Agent — Streamlit Application
Finance Department
Diagnostic reporting for QuickBooks month-end close.

Reports:
  1. Flux (Variance) Narrative
  2. Recurring Vendor Gap Analysis ("Missing Bill")
  3. Suspense & Misc Resolution Worksheet
  4. Materiality & Risk Threshold
  5. IIF Import Pre-Flight Validation
  6. Multi-Source Reconciliation Summary
  7. Top 20 Balance Sheet Accounts — Closing Month vs Preceding Months Avg
  8. Top 20 Profit & Loss Accounts — Closing Month vs Preceding Months Avg
"""

import streamlit as st
import pandas as pd
import numpy as np
import re
import io
import os
import subprocess
import tempfile
import json
import datetime
import warnings
from collections import defaultdict

warnings.filterwarnings("ignore")

# ─────────────────────────────────────────────────────────────
# 0 · APPLE-BRANDED THEME
# ─────────────────────────────────────────────────────────────
APPLE = {
    "blue":        "#0071E3",
    "deep_blue":   "#0055CC",
    "bright_blue": "#1A8CFF",
    "near_black":  "#1D1D1F",
    "gray_sec":    "#6E6E73",
    "sys_gray":    "#8E8E93",
    "mid_gray":    "#D2D2D7",
    "light_gray":  "#F5F5F7",
    "white":       "#FFFFFF",
    "red":         "#FF3B30",
    "green":       "#34C759",
    "orange":      "#FF9500",
    "yellow":      "#FFCC00",
}

st.set_page_config(
    page_title="Month-End Close Agent · Month-End Close",
    page_icon="🏢",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown(f"""
<style>
    /* ── Global ── */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
    html, body, [class*="css"] {{
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Helvetica Neue', sans-serif;
        color: {APPLE["near_black"]};
    }}
    /* ── Header banner ── */
    .apple-header {{
        background: linear-gradient(135deg, {APPLE["bright_blue"]}, {APPLE["deep_blue"]});
        border-radius: 16px;
        padding: 2rem 2.5rem;
        margin-bottom: 1.5rem;
        color: {APPLE["white"]};
    }}
    .apple-header h1 {{
        font-weight: 700; font-size: 1.75rem; margin: 0; letter-spacing: -0.02em;
    }}
    .apple-header p {{
        font-weight: 400; font-size: 0.95rem; opacity: 0.85; margin: 0.35rem 0 0 0;
    }}
    /* ── Cards ── */
    .metric-card {{
        background: {APPLE["light_gray"]};
        border: 1px solid {APPLE["mid_gray"]};
        border-radius: 12px;
        padding: 1.15rem 1.35rem;
        margin-bottom: 0.75rem;
    }}
    .metric-card .label {{
        font-size: 0.75rem; font-weight: 500; color: {APPLE["gray_sec"]};
        text-transform: uppercase; letter-spacing: 0.04em;
    }}
    .metric-card .value {{
        font-size: 1.5rem; font-weight: 700; color: {APPLE["near_black"]};
    }}
    /* ── Status pills ── */
    .pill-pass {{
        background: {APPLE["green"]}22; color: {APPLE["green"]};
        padding: 2px 10px; border-radius: 20px; font-weight: 600; font-size: 0.8rem;
    }}
    .pill-fail {{
        background: {APPLE["red"]}22; color: {APPLE["red"]};
        padding: 2px 10px; border-radius: 20px; font-weight: 600; font-size: 0.8rem;
    }}
    .pill-warn {{
        background: {APPLE["orange"]}22; color: {APPLE["orange"]};
        padding: 2px 10px; border-radius: 20px; font-weight: 600; font-size: 0.8rem;
    }}
    /* ── Section headers ── */
    .section-head {{
        font-weight: 700; font-size: 1.05rem; color: {APPLE["near_black"]};
        border-bottom: 2px solid {APPLE["blue"]};
        padding-bottom: 0.4rem; margin: 1.5rem 0 0.75rem 0;
        text-transform: uppercase; letter-spacing: 0.03em;
    }}
    /* ── Table tweaks ── */
    .stDataFrame thead th {{
        background: {APPLE["light_gray"]} !important;
        font-weight: 600 !important;
        color: {APPLE["near_black"]} !important;
    }}
    /* ── Sidebar ── */
    section[data-testid="stSidebar"] {{
        background: {APPLE["light_gray"]};
    }}
    section[data-testid="stSidebar"] .stMarkdown p {{
        color: {APPLE["gray_sec"]};
        font-size: 0.85rem;
    }}
    /* ── Narrative blocks ── */
    .narrative-block {{
        background: {APPLE["white"]};
        border-left: 4px solid {APPLE["blue"]};
        padding: 1rem 1.25rem;
        border-radius: 0 8px 8px 0;
        margin-bottom: 0.75rem;
        font-size: 0.92rem;
        line-height: 1.6;
    }}
    /* ── Footer ── */
    .apple-footer {{
        text-align: center;
        color: {APPLE["sys_gray"]};
        font-size: 0.75rem;
        margin-top: 3rem;
        padding-top: 1rem;
        border-top: 1px solid {APPLE["mid_gray"]};
    }}
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────
# 1 · HELPER FUNCTIONS
# ─────────────────────────────────────────────────────────────

def metric_card(label: str, value: str) -> str:
    return f"""<div class="metric-card">
        <div class="label">{label}</div>
        <div class="value">{value}</div>
    </div>"""


def pill(status: str, text: str) -> str:
    cls = {"pass": "pill-pass", "fail": "pill-fail", "warn": "pill-warn"}.get(status, "pill-warn")
    return f'<span class="{cls}">{text}</span>'


def section(title: str):
    st.markdown(f'<div class="section-head">{title}</div>', unsafe_allow_html=True)


def narrative(text: str):
    st.markdown(f'<div class="narrative-block">{text}</div>', unsafe_allow_html=True)


def fmt_currency(v):
    try:
        return f"${abs(float(v)):,.2f}"
    except (ValueError, TypeError):
        return "$0.00"


# ─────────────────────────────────────────────────────────────
# 2 · DATA PARSERS
# ─────────────────────────────────────────────────────────────

@st.cache_data
def parse_gl(file) -> pd.DataFrame:
    """Parse a General Ledger from CSV or QuickBooks Excel export.

    Sniffs the actual file content to detect Excel binaries even when
    the extension is .csv (common with QuickBooks exports).
    """
    # Read first 4 bytes to detect file type
    header = file.read(4)
    file.seek(0)
    is_excel = header[:2] == b"PK" or header[:8] == b"\xd0\xcf\x11\xe0"

    name = file.name.lower()
    if is_excel or name.endswith((".xlsx", ".xls")):
        return _parse_gl_excel(file)
    else:
        return _parse_gl_csv(file)


def _parse_gl_csv(file) -> pd.DataFrame:
    """Handles standard CSV GL exports with named columns."""
    # Try multiple encodings for robustness
    for encoding in ["utf-8", "latin-1", "cp1252"]:
        try:
            file.seek(0)
            df = pd.read_csv(file, encoding=encoding)
            break
        except (UnicodeDecodeError, pd.errors.ParserError):
            continue
    else:
        file.seek(0)
        df = pd.read_csv(file, encoding="latin-1", on_bad_lines="skip")
    df.columns = [c.strip() for c in df.columns]
    col_map = {}
    for c in df.columns:
        cl = c.lower()
        if "date" in cl:
            col_map["Date"] = c
        elif "account" in cl and "code" in cl:
            col_map["Account"] = c
        elif "account" in cl and "Account" not in col_map:
            col_map["Account"] = c
        elif "name" in cl or "vendor" in cl:
            col_map["Name"] = c
        elif "memo" in cl or "desc" in cl:
            col_map["Memo"] = c
        elif "debit" in cl:
            col_map["Debit"] = c
        elif "credit" in cl:
            col_map["Credit"] = c
        elif "type" in cl:
            col_map["Type"] = c
        elif "split" in cl:
            col_map["Split"] = c
        elif "num" in cl:
            col_map["Num"] = c
    rename = {v: k for k, v in col_map.items()}
    df = df.rename(columns=rename)
    return _clean_gl(df)


def _parse_gl_excel(file) -> pd.DataFrame:
    """Parse QuickBooks Desktop hierarchical Excel GL export."""
    xls = pd.ExcelFile(file, engine="openpyxl")
    sheets = xls.sheet_names
    target = "Sheet1" if "Sheet1" in sheets else sheets[-1]
    raw = pd.read_excel(file, sheet_name=target, engine="openpyxl", header=None)

    # Detect header row — look for a row containing "Type" and "Date"
    header_idx = None
    for i in range(min(10, len(raw))):
        row_vals = raw.iloc[i].astype(str).str.strip().str.lower().tolist()
        if "type" in row_vals and "date" in row_vals:
            header_idx = i
            break
    if header_idx is None:
        header_idx = 0

    col_indices = {}
    known_header_names = {"type", "date", "num", "name", "memo", "split", "debit", "credit", "balance"}
    for idx in range(raw.shape[1]):
        val = raw.iloc[header_idx, idx]
        if pd.isna(val):
            continue
        vl = str(val).lower().strip()
        if vl in known_header_names:
            col_indices[vl.capitalize()] = idx

    # Determine account columns: columns BEFORE the first known field column.
    # In QuickBooks exports, account hierarchy is in the leftmost columns,
    # and data fields (Type, Date, etc.) start after that.
    first_field_col = min(col_indices.values()) if col_indices else 0
    acct_cols = [c for c in range(first_field_col)]

    # Build a hierarchy tracker — deepest non-null column wins
    records = []
    # Track the most recent account value at each column level
    acct_stack = {}
    current_account = ""

    for i in range(header_idx + 1, len(raw)):
        row = raw.iloc[i]
        # Transaction rows have a date
        date_val = row.get(col_indices.get("Date"), None)

        # First, check if any account column has a value to update the hierarchy
        for ac in acct_cols:
            v = row.get(ac)
            if pd.notna(v):
                val_str = str(v).strip()
                if val_str and not val_str.lower().startswith("total"):
                    acct_stack[ac] = val_str
                    # Clear deeper levels when a higher level changes
                    for deeper in acct_cols:
                        if deeper > ac:
                            acct_stack.pop(deeper, None)

        if pd.notna(date_val) and str(date_val).strip() != "":
            # Use the deepest level in the hierarchy stack
            if acct_stack:
                deepest_col = max(acct_stack.keys())
                current_account = acct_stack[deepest_col]
            rec = {"Account": current_account}
            for field, cidx in col_indices.items():
                rec[field] = row.get(cidx)
            records.append(rec)

    df = pd.DataFrame(records)
    return _clean_gl(df)


def _clean_gl(df: pd.DataFrame) -> pd.DataFrame:
    """Normalize GL DataFrame."""
    if "Date" in df.columns:
        df["Date"] = pd.to_datetime(df["Date"], errors="coerce")
    for col in ["Debit", "Credit"]:
        if col in df.columns:
            df[col] = pd.to_numeric(
                df[col].astype(str).str.replace(r"[,$]", "", regex=True),
                errors="coerce",
            ).fillna(0.0)
    for col in ["Name", "Memo", "Account", "Type", "Split"]:
        if col in df.columns:
            df[col] = df[col].astype(str).fillna("")
    # Add helper columns
    df["Amount"] = df.get("Debit", 0) - df.get("Credit", 0)
    if "Date" in df.columns:
        df["YearMonth"] = df["Date"].dt.to_period("M")
    # Strip account number/name
    if "Account" in df.columns:
        df["AcctClean"] = df["Account"].str.replace(r"^\d+\s*·\s*", "", regex=True).str.strip()
    return df


@st.cache_data
def parse_iif(file) -> dict:
    """Parse a QuickBooks IIF file into accounts and classes."""
    content = file.read().decode("utf-8", errors="replace")
    lines = content.split("\n")
    accounts, classes = [], []
    section = None
    headers = []
    for line in lines:
        parts = line.rstrip("\r").split("\t")
        if not parts:
            continue
        tag = parts[0].strip()
        if tag == "!ACCNT":
            section = "ACCNT"
            headers = [p.strip() for p in parts]
        elif tag == "!CLASS":
            section = "CLASS"
            headers = [p.strip() for p in parts]
        elif tag == "ACCNT" and section == "ACCNT":
            row = dict(zip(headers, [p.strip() for p in parts]))
            accounts.append(row)
        elif tag == "CLASS" and section == "CLASS":
            row = dict(zip(headers, [p.strip() for p in parts]))
            classes.append(row)
        elif tag.startswith("!") and tag not in ("!HDR",):
            section = None

    acct_df = pd.DataFrame(accounts) if accounts else pd.DataFrame(columns=["NAME", "ACCNTTYPE", "ACCNUM", "HIDDEN"])
    class_df = pd.DataFrame(classes) if classes else pd.DataFrame(columns=["NAME", "HIDDEN"])
    return {
        "accounts": acct_df,
        "classes": class_df,
        "account_names": set(acct_df["NAME"].tolist()) if "NAME" in acct_df.columns else set(),
        "active_accounts": set(
            acct_df.loc[acct_df.get("HIDDEN", pd.Series(["N"] * len(acct_df))) == "N", "NAME"].tolist()
        ) if "NAME" in acct_df.columns else set(),
        "active_classes": set(
            class_df.loc[class_df.get("HIDDEN", pd.Series(["N"] * len(class_df))) == "N", "NAME"].tolist()
        ) if "NAME" in class_df.columns else set(),
    }


@st.cache_data
def parse_coa_csv(file) -> dict:
    """Parse a CSV-based Chart of Accounts."""
    df = pd.read_csv(file)
    df.columns = [c.strip() for c in df.columns]
    names = set()
    for c in df.columns:
        if "name" in c.lower() or "account" in c.lower():
            names.update(df[c].dropna().astype(str).tolist())
    return {
        "accounts": df,
        "classes": pd.DataFrame(),
        "account_names": names,
        "active_accounts": names,
        "active_classes": set(),
    }


def extract_pdf_text(file) -> str:
    """Best-effort text extraction from an uploaded PDF."""
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file.read())) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)
    except ImportError:
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(file.read()))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception:
            return ""
    except Exception:
        return ""


# ─────────────────────────────────────────────────────────────
# 2b · BALANCE REPORT PARSER (Monthly BS / P&L)
# ─────────────────────────────────────────────────────────────

_MONTH_ABBREVS = {
    "jan": 1, "feb": 2, "mar": 3, "apr": 4, "may": 5, "jun": 6,
    "jul": 7, "aug": 8, "sep": 9, "oct": 10, "nov": 11, "dec": 12,
    "january": 1, "february": 2, "march": 3, "april": 4,
    "june": 6, "july": 7, "august": 8, "september": 9,
    "october": 10, "november": 11, "december": 12,
}


def _detect_month_from_header(header_str: str):
    """Try to extract a (year, month) from a column header like 'Jul 2025', 'July 25', '2025-07', etc.

    Returns (year, month) tuple or None.
    """
    s = str(header_str).strip()
    if not s or s.lower() in ("nan", "none", "total", ""):
        return None

    # Pattern: "Jul 2025", "July 2025", "Jul 25", "July '25"
    m = re.match(r"(?i)([a-z]+)['\s\-,.]*(\d{2,4})", s)
    if m:
        mon_str = m.group(1).lower()
        yr_str = m.group(2)
        mon = _MONTH_ABBREVS.get(mon_str)
        if mon:
            yr = int(yr_str)
            if yr < 100:
                yr += 2000
            return (yr, mon)

    # Pattern: "2025-07", "2025/07"
    m = re.match(r"(\d{4})[\-/](\d{1,2})", s)
    if m:
        return (int(m.group(1)), int(m.group(2)))

    # Pattern: "07/2025", "7/2025"
    m = re.match(r"(\d{1,2})[\-/](\d{4})", s)
    if m:
        mon, yr = int(m.group(1)), int(m.group(2))
        if 1 <= mon <= 12:
            return (yr, mon)

    return None


def parse_balance_report(file) -> pd.DataFrame | None:
    """Parse a monthly balance report (BS or P&L) where accounts are rows and months are columns.

    Returns a DataFrame with columns: Account, and one column per detected month
    labeled as 'YYYY-MM' (e.g., '2025-07'), values are floats.
    Returns None if parsing fails.
    """
    try:
        header = file.read(4)
        file.seek(0)
        is_excel = header[:2] == b"PK" or header[:4] == b"\xd0\xcf\x11\xe0"

        name = file.name.lower()
        if is_excel or name.endswith((".xlsx", ".xls")):
            raw = pd.read_excel(file, header=None, engine="openpyxl")
        else:
            file.seek(0)
            raw = pd.read_csv(file, header=None, encoding="latin-1")
    except Exception:
        return None

    if raw.empty or raw.shape[1] < 2:
        return None

    # Detect the header row — the one with the most month-parseable columns
    best_row = 0
    best_count = 0
    best_map = {}
    for i in range(min(15, len(raw))):
        month_map = {}
        for ci in range(raw.shape[1]):
            val = raw.iloc[i, ci]
            parsed = _detect_month_from_header(val)
            if parsed:
                month_map[ci] = parsed
        if len(month_map) > best_count:
            best_count = len(month_map)
            best_row = i
            best_map = month_map

    if best_count < 1:
        return None

    # Determine the account column (usually the first text column before the month columns)
    first_month_col = min(best_map.keys())
    acct_col = 0
    for ci in range(first_month_col):
        # Pick the column with the most non-empty string values
        vals = raw.iloc[best_row + 1:, ci].dropna()
        if len(vals) > 0:
            acct_col = ci
            break

    # Build clean DataFrame
    records = []
    for i in range(best_row + 1, len(raw)):
        acct_val = raw.iloc[i, acct_col]
        if pd.isna(acct_val) or str(acct_val).strip() == "":
            continue
        acct_name = str(acct_val).strip()
        # Skip total/header rows
        if re.match(r"(?i)^(total|net\s|other\s*(total|net)|^$)", acct_name):
            continue

        row = {"Account": acct_name}
        has_value = False
        for ci, (yr, mon) in best_map.items():
            cell = raw.iloc[i, ci]
            # Parse numeric value, stripping $ , ( ) etc.
            if pd.isna(cell):
                num = 0.0
            else:
                cleaned = re.sub(r"[,$\s]", "", str(cell))
                # Handle parenthetical negatives
                neg = cleaned.startswith("(") and cleaned.endswith(")")
                cleaned = cleaned.strip("()")
                try:
                    num = float(cleaned)
                    if neg:
                        num = -num
                    has_value = True
                except ValueError:
                    num = 0.0
            col_label = f"{yr}-{mon:02d}"
            row[col_label] = num

        if has_value:
            records.append(row)

    if not records:
        return None

    df = pd.DataFrame(records)
    # Consolidate duplicate account names (sum them)
    month_cols = [c for c in df.columns if c != "Account"]
    if month_cols:
        df = df.groupby("Account", as_index=False)[month_cols].sum()
    return df


# ─────────────────────────────────────────────────────────────
# 3 · REPORT ENGINES
# ─────────────────────────────────────────────────────────────

def report_flux(gl: pd.DataFrame):
    """Report 1 — Flux (Variance) Narrative."""
    section("Flux (Variance) Narrative Report")

    if "YearMonth" not in gl.columns or gl["YearMonth"].isna().all():
        st.warning("Date column could not be parsed — unable to compute month-over-month flux.")
        return

    periods = sorted(gl["YearMonth"].dropna().unique())
    if len(periods) < 2:
        st.info("Only one period detected in the GL. At least two months are required for variance analysis.")
        return

    curr = periods[-1]
    prev = periods[-2]
    st.markdown(f"**Comparing:** `{prev}` → `{curr}`")

    curr_df = gl[gl["YearMonth"] == curr]
    prev_df = gl[gl["YearMonth"] == prev]

    curr_totals = curr_df.groupby("Account")["Amount"].sum()
    prev_totals = prev_df.groupby("Account")["Amount"].sum()

    flux = pd.DataFrame({"Current": curr_totals, "Prior": prev_totals}).fillna(0)
    flux["Variance"] = flux["Current"] - flux["Prior"]
    flux["Var_%"] = np.where(flux["Prior"] != 0, (flux["Variance"] / flux["Prior"].abs()) * 100, np.nan)
    flux = flux.sort_values("Variance", key=abs, ascending=False)

    # Summary metrics
    c1, c2, c3 = st.columns(3)
    total_var = flux["Variance"].sum()
    top_driver = flux.index[0] if len(flux) > 0 else "N/A"
    large_moves = (flux["Variance"].abs() > flux["Variance"].abs().quantile(0.9)).sum()
    with c1:
        st.markdown(metric_card("Net Variance", f"${total_var:,.2f}"), unsafe_allow_html=True)
    with c2:
        st.markdown(metric_card("Top Driver", str(top_driver)[:40]), unsafe_allow_html=True)
    with c3:
        st.markdown(metric_card("Large Moves (>P90)", str(large_moves)), unsafe_allow_html=True)

    # Narrative generation — scan memos and vendors for top movers
    section("AI-Generated Narrative")
    top_n = flux.head(10)
    for acct, row in top_n.iterrows():
        direction = "increased" if row["Variance"] > 0 else "decreased"
        pct = f"{abs(row['Var_%']):.1f}%" if pd.notna(row["Var_%"]) else "N/A"
        # Scan memos for this account in current period
        acct_txns = curr_df[curr_df["Account"] == acct]
        memo_sample = acct_txns["Memo"].dropna().unique()[:3]
        vendor_sample = acct_txns["Name"].dropna().unique()[:3]
        memo_str = "; ".join(str(m) for m in memo_sample if str(m).strip() and str(m) != "nan") or "no memo detail"
        vendor_str = ", ".join(str(v) for v in vendor_sample if str(v).strip() and str(v) != "nan") or "various"

        # Detect patterns
        notes = []
        if acct_txns.shape[0] > 0:
            avg_txn = acct_txns["Amount"].abs().mean()
            max_txn = acct_txns["Amount"].abs().max()
            if max_txn > avg_txn * 3 and acct_txns.shape[0] > 1:
                notes.append("an unusually large single transaction was detected")
            dup_memos = acct_txns.groupby("Memo").size()
            dup_memos = dup_memos[dup_memos > 1]
            if len(dup_memos) > 0:
                notes.append(f"possible duplicate entries for: {', '.join(dup_memos.index[:2])}")

        note_str = ". Additionally, " + "; ".join(notes) + "." if notes else "."
        txt = (
            f"<strong>{acct}</strong> {direction} by <strong>${abs(row['Variance']):,.2f}</strong> "
            f"({pct}), moving from ${row['Prior']:,.2f} to ${row['Current']:,.2f}. "
            f"Key vendors: {vendor_str}. Recent memos reference: <em>{memo_str}</em>{note_str}"
        )
        narrative(txt)

    # Detail table
    section("Variance Detail")
    display = flux.copy()
    display.index.name = "Account"
    display = display.reset_index()
    for c in ["Current", "Prior", "Variance"]:
        display[c] = display[c].map(lambda v: f"${v:,.2f}")
    display["Var_%"] = display["Var_%"].map(lambda v: f"{v:.1f}%" if pd.notna(v) else "—")
    st.dataframe(display, use_container_width=True, hide_index=True)


def report_vendor_gap(gl: pd.DataFrame, closing_month: int = None, start_month: int = None):
    """Report 2 — Recurring Vendor Gap Analysis ('Missing Bill')."""
    section("Recurring Vendor Gap Analysis")

    if "Name" not in gl.columns or "YearMonth" not in gl.columns:
        st.warning("GL must contain Name and Date columns for vendor gap analysis.")
        return

    vendors = gl[gl["Name"].str.strip() != "nan"].copy()
    if vendors.empty:
        st.info("No vendor names found in the GL.")
        return

    all_periods = sorted(vendors["YearMonth"].dropna().unique())
    if len(all_periods) < 3:
        st.info("At least 3 months of history are needed to detect recurring vendors.")
        return

    # Determine closing period and preceding periods from user selection
    if closing_month and start_month:
        closing_periods = [p for p in all_periods if p.month == closing_month]
        prec_month_nums = _get_preceding_months(closing_month, start_month)
        preceding_periods = sorted([p for p in all_periods if p.month in prec_month_nums])
        close_p = closing_periods[-1] if closing_periods else all_periods[-1]
        analysis_periods = preceding_periods + [close_p]
        # Filter to relevant periods
        vendors_in_scope = vendors[vendors["YearMonth"].isin(analysis_periods)]
    else:
        close_p = all_periods[-1]
        preceding_periods = all_periods[:-1]
        analysis_periods = all_periods
        vendors_in_scope = vendors

    curr = close_p
    history = vendors_in_scope[vendors_in_scope["YearMonth"] != curr]
    curr_df = vendors_in_scope[vendors_in_scope["YearMonth"] == curr]
    curr_vendors = set(curr_df["Name"].unique())

    vendor_periods = history.groupby("Name")["YearMonth"].nunique()
    total_hist_months = len(preceding_periods)
    recurring = vendor_periods[vendor_periods >= max(2, total_hist_months * 0.5)]

    missing = [v for v in recurring.index if v not in curr_vendors and str(v).strip() and str(v) != "nan"]

    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown(metric_card("Recurring Vendors", str(len(recurring))), unsafe_allow_html=True)
    with c2:
        st.markdown(metric_card("Missing This Month", str(len(missing))), unsafe_allow_html=True)
    with c3:
        est_accrual = 0
        for v in missing:
            avg = history[history["Name"] == v]["Credit"].mean()
            est_accrual += avg if pd.notna(avg) else 0
        st.markdown(metric_card("Est. Accrual Total", f"${est_accrual:,.2f}"), unsafe_allow_html=True)

    if missing:
        section("Suggested Accruals")
        rows = []
        for v in missing:
            v_hist = history[history["Name"] == v]
            avg_amt = v_hist["Credit"].mean()
            last_date = v_hist["Date"].max()
            freq = v_hist["YearMonth"].nunique()
            typical_acct = v_hist["Account"].mode().iloc[0] if not v_hist["Account"].mode().empty else "Unknown"
            rows.append({
                "Vendor": v,
                "Historical Frequency": f"{freq} of {total_hist_months} months",
                "Avg Monthly Amount": f"${avg_amt:,.2f}" if pd.notna(avg_amt) else "—",
                "Last Seen": str(last_date.date()) if pd.notna(last_date) else "—",
                "Typical Account": typical_acct,
                "Suggested Accrual": f"${avg_amt:,.2f}" if pd.notna(avg_amt) else "—",
            })
        st.dataframe(pd.DataFrame(rows), use_container_width=True, hide_index=True)

        # ── Per-month transaction breakdown ──
        section("Monthly Transaction Detail by Vendor")
        st.markdown(f"Transactions per month for missing vendors across "
                     f"**{', '.join(str(p) for p in preceding_periods)}** and closing month **{curr}**.")

        month_cols = [str(p) for p in analysis_periods]
        detail_rows = []
        for v in missing:
            v_all = vendors_in_scope[vendors_in_scope["Name"] == v]
            row = {"Vendor": v}
            for p in analysis_periods:
                p_txns = v_all[v_all["YearMonth"] == p]
                p_total = p_txns["Credit"].sum()
                row[str(p)] = p_total if p_total != 0 else 0
            detail_rows.append(row)

        detail_df = pd.DataFrame(detail_rows)
        # Format for display
        disp = detail_df.copy()
        for mc in month_cols:
            if mc in disp.columns:
                disp[mc] = disp[mc].map(lambda v: f"${v:,.2f}" if isinstance(v, (int, float)) and v != 0 else "—")
        st.dataframe(disp, use_container_width=True, hide_index=True)
    else:
        st.success("All recurring vendors have activity in the current period.")


def report_suspense(gl: pd.DataFrame, coa: dict | None):
    """Report 3 — Suspense & Misc Resolution Worksheet."""
    section("Suspense & Misc Resolution Worksheet")

    patterns = r"(?i)(suspense|misc|other|unclass|uncategoriz|clearing|unknown|unallocat)"
    mask = gl["Account"].str.contains(patterns, na=False) | gl["Memo"].str.contains(patterns, na=False)
    if "Split" in gl.columns:
        mask = mask | gl["Split"].str.contains(patterns, na=False)

    flagged = gl[mask].copy()

    c1, c2 = st.columns(2)
    with c1:
        st.markdown(metric_card("Flagged Transactions", str(len(flagged))), unsafe_allow_html=True)
    with c2:
        total_amt = flagged["Amount"].abs().sum()
        st.markdown(metric_card("Total at Risk", f"${total_amt:,.2f}"), unsafe_allow_html=True)

    if flagged.empty:
        st.success("No transactions found in suspense, misc, or clearing accounts.")
        return

    # Suggest reclassifications using COA keyword matching
    suggestions = []
    active_accts = coa.get("active_accounts", set()) if coa else set()
    acct_list = sorted(active_accts)

    for _, txn in flagged.iterrows():
        memo = str(txn.get("Memo", "")).lower()
        name = str(txn.get("Name", "")).lower()
        search_text = memo + " " + name

        best_match = ""
        best_score = 0
        for acct in acct_list:
            acct_lower = acct.lower()
            # Split account name into keywords
            keywords = re.split(r"[:\s\-&/]+", acct_lower)
            keywords = [k for k in keywords if len(k) > 2]
            score = sum(1 for k in keywords if k in search_text)
            if score > best_score:
                best_score = score
                best_match = acct

        suggestions.append({
            "Date": str(txn.get("Date", ""))[:10],
            "Current Account": txn.get("Account", ""),
            "Name / Vendor": txn.get("Name", ""),
            "Memo": str(txn.get("Memo", ""))[:80],
            "Amount": f"${txn['Amount']:,.2f}",
            "Suggested Reclass": best_match if best_score >= 1 else "— manual review —",
            "Confidence": f"{'High' if best_score >= 3 else 'Medium' if best_score >= 1 else 'Low'}",
        })

    st.dataframe(pd.DataFrame(suggestions), use_container_width=True, hide_index=True)


def report_materiality(gl: pd.DataFrame, threshold: float):
    """Report 4 — Materiality & Risk Threshold."""
    section("Materiality & Risk Threshold Report")

    st.markdown(f"**Active Threshold:** ${threshold:,.0f}")

    # Flag categories: large uncategorized, suspense, large single txns, imbalances
    flags = []

    # Large transactions above threshold
    large = gl[gl["Amount"].abs() >= threshold].copy()
    for _, txn in large.iterrows():
        risk = "Low"
        reasons = []

        acct = str(txn.get("Account", "")).lower()
        memo = str(txn.get("Memo", "")).lower()

        if re.search(r"(suspense|misc|other|clearing|unknown)", acct):
            risk = "High"
            reasons.append("sitting in suspense/misc account")
        if memo in ("", "nan", "none"):
            risk = "High" if risk == "High" else "Medium"
            reasons.append("no memo/description")
        if abs(txn["Amount"]) >= threshold * 5:
            risk = "High"
            reasons.append("exceeds 5x materiality threshold")

        flags.append({
            "Date": str(txn.get("Date", ""))[:10],
            "Account": txn.get("Account", ""),
            "Name": txn.get("Name", ""),
            "Memo": str(txn.get("Memo", ""))[:60],
            "Amount": f"${txn['Amount']:,.2f}",
            "Risk": risk,
            "Reason": "; ".join(reasons) if reasons else "Material amount",
        })

    df_flags = pd.DataFrame(flags)
    if df_flags.empty:
        st.success(f"No transactions exceed the ${threshold:,.0f} materiality threshold.")
        return

    # Summary
    c1, c2, c3 = st.columns(3)
    high_ct = (df_flags["Risk"] == "High").sum()
    med_ct = (df_flags["Risk"] == "Medium").sum()
    with c1:
        st.markdown(metric_card("Flagged Items", str(len(df_flags))), unsafe_allow_html=True)
    with c2:
        st.markdown(metric_card("High Risk", f"{pill('fail', str(high_ct))}"), unsafe_allow_html=True)
    with c3:
        st.markdown(metric_card("Medium Risk", f"{pill('warn', str(med_ct))}"), unsafe_allow_html=True)

    st.dataframe(df_flags.sort_values("Risk", ascending=True), use_container_width=True, hide_index=True)


def report_iif_preflight(gl: pd.DataFrame, coa: dict | None):
    """Report 5 — IIF Import Pre-Flight Validation."""
    section("IIF Import Pre-Flight Validation")

    checks = []

    # Check 1: Debits = Credits per period
    if "YearMonth" in gl.columns:
        for period in sorted(gl["YearMonth"].dropna().unique()):
            p_df = gl[gl["YearMonth"] == period]
            total_dr = p_df["Debit"].sum()
            total_cr = p_df["Credit"].sum()
            diff = abs(total_dr - total_cr)
            balanced = diff < 0.01
            checks.append({
                "Check": f"Debits = Credits ({period})",
                "Status": "PASS" if balanced else "FAIL",
                "Detail": f"DR: ${total_dr:,.2f}  |  CR: ${total_cr:,.2f}  |  Diff: ${diff:,.2f}",
            })
    else:
        total_dr = gl["Debit"].sum()
        total_cr = gl["Credit"].sum()
        diff = abs(total_dr - total_cr)
        checks.append({
            "Check": "Debits = Credits (all data)",
            "Status": "PASS" if diff < 0.01 else "FAIL",
            "Detail": f"DR: ${total_dr:,.2f}  |  CR: ${total_cr:,.2f}  |  Diff: ${diff:,.2f}",
        })

    # Check 2: Account names match COA
    if coa:
        gl_accounts = set(gl["Account"].unique())
        active = coa.get("active_accounts", set())
        all_known = coa.get("account_names", set())
        unmatched = []
        for a in gl_accounts:
            a_clean = a.strip()
            if not a_clean or a_clean == "nan":
                continue
            # Fuzzy: check if GL account name is a substring of any COA name or vice versa
            found = False
            for known in all_known:
                if a_clean == known or a_clean in known or known in a_clean:
                    found = True
                    break
            if not found:
                # Try matching just the account number portion
                num_match = re.match(r"^(\d+)", a_clean)
                if num_match:
                    num = num_match.group(1)
                    for known in all_known:
                        if num in known:
                            found = True
                            break
            if not found:
                unmatched.append(a_clean)

        checks.append({
            "Check": "All GL accounts exist in COA",
            "Status": "PASS" if len(unmatched) == 0 else "WARN",
            "Detail": f"{len(unmatched)} unmatched account(s)" + (f": {', '.join(unmatched[:5])}" if unmatched else ""),
        })

        # Check 3: Active classes
        active_cls = coa.get("active_classes", set())
        if active_cls:
            checks.append({
                "Check": "Active classes loaded",
                "Status": "PASS",
                "Detail": f"{len(active_cls)} active class(es) available for validation",
            })
    else:
        checks.append({
            "Check": "COA validation",
            "Status": "SKIP",
            "Detail": "No Chart of Accounts uploaded — skipping account name validation",
        })

    # Check 4: No blank accounts
    blank_accts = gl["Account"].isna().sum() + (gl["Account"] == "").sum() + (gl["Account"] == "nan").sum()
    checks.append({
        "Check": "No blank account codes",
        "Status": "PASS" if blank_accts == 0 else "WARN",
        "Detail": f"{blank_accts} transaction(s) with blank account" if blank_accts else "All transactions have accounts",
    })

    # Check 5: No future dates
    if "Date" in gl.columns:
        future = gl[gl["Date"] > pd.Timestamp.now()].shape[0]
        checks.append({
            "Check": "No future-dated transactions",
            "Status": "PASS" if future == 0 else "WARN",
            "Detail": f"{future} future-dated transaction(s)" if future else "All dates are current or past",
        })

    # Display
    for ck in checks:
        status = ck["Status"]
        if status == "PASS":
            icon = pill("pass", "PASS")
        elif status == "FAIL":
            icon = pill("fail", "FAIL")
        else:
            icon = pill("warn", status)
        st.markdown(f"{icon} &nbsp; **{ck['Check']}** — {ck['Detail']}", unsafe_allow_html=True)

    return checks


def report_reconciliation(gl: pd.DataFrame, pdf_texts: list[str]):
    """Report 6 — Multi-Source Reconciliation Summary."""
    section("Multi-Source Reconciliation Summary")

    if not pdf_texts:
        st.info(
            "Upload bank statements or invoices (PDF) in the sidebar to enable "
            "three-way reconciliation. The system will attempt to match GL entries "
            "against extracted PDF line items."
        )

    # Bank-side: transactions with bank-related accounts
    bank_kw = r"(?i)(bank|cash|checking|savings|operating|deposit)"
    bank_gl = gl[gl["Account"].str.contains(bank_kw, na=False)].copy()

    c1, c2 = st.columns(2)
    with c1:
        st.markdown(metric_card("GL Bank Transactions", str(len(bank_gl))), unsafe_allow_html=True)
    with c2:
        st.markdown(metric_card("PDF Documents", str(len(pdf_texts))), unsafe_allow_html=True)

    # Extract amounts from PDF text
    if pdf_texts:
        section("PDF-Extracted Line Items")
        all_amounts = []
        for idx, text in enumerate(pdf_texts):
            # Find dollar amounts in the text
            amounts = re.findall(r"\$?([\d,]+\.\d{2})", text)
            for a in amounts:
                val = float(a.replace(",", ""))
                if val > 0.01:
                    all_amounts.append({"Source": f"PDF #{idx+1}", "Amount": val})

        if all_amounts:
            pdf_df = pd.DataFrame(all_amounts)
            st.dataframe(pdf_df.head(50), use_container_width=True, hide_index=True)

            # Match against GL
            section("Reconciliation Matches")
            matches = []
            unmatched_pdf = []
            gl_amounts = bank_gl[["Date", "Account", "Memo", "Debit", "Credit"]].copy()

            for pa in all_amounts:
                val = pa["Amount"]
                # Try matching debit or credit
                match_dr = gl_amounts[(gl_amounts["Debit"] - val).abs() < 0.01]
                match_cr = gl_amounts[(gl_amounts["Credit"] - val).abs() < 0.01]
                if not match_dr.empty:
                    r = match_dr.iloc[0]
                    matches.append({
                        "PDF Amount": f"${val:,.2f}",
                        "GL Date": str(r["Date"])[:10],
                        "GL Account": r["Account"],
                        "GL Memo": str(r["Memo"])[:50],
                        "GL Amount (DR)": f"${r['Debit']:,.2f}",
                        "Status": "Matched",
                    })
                elif not match_cr.empty:
                    r = match_cr.iloc[0]
                    matches.append({
                        "PDF Amount": f"${val:,.2f}",
                        "GL Date": str(r["Date"])[:10],
                        "GL Account": r["Account"],
                        "GL Memo": str(r["Memo"])[:50],
                        "GL Amount (CR)": f"${r['Credit']:,.2f}",
                        "Status": "Matched",
                    })
                else:
                    unmatched_pdf.append({"PDF Amount": f"${val:,.2f}", "Status": "Unmatched"})

            if matches:
                st.dataframe(pd.DataFrame(matches), use_container_width=True, hide_index=True)
            if unmatched_pdf:
                st.markdown(f"**{len(unmatched_pdf)} PDF amount(s) unmatched** in the GL — these may represent "
                            "sales tax, shipping, or fees that banking rules missed.")
        else:
            st.warning("Could not extract numeric amounts from the uploaded PDF(s).")

    # Sales tax & shipping scan
    section("Sales Tax & Shipping Scan")
    tax_kw = r"(?i)(sales\s*tax|tax|shipping|freight|delivery|handling|surcharge)"
    tax_txns = gl[gl["Memo"].str.contains(tax_kw, na=False) | gl["Account"].str.contains(tax_kw, na=False)]
    if tax_txns.empty:
        st.info("No explicit sales tax or shipping entries found in the GL. "
                "These are common items that bank rules often miss — verify manually.")
    else:
        display = tax_txns[["Date", "Account", "Name", "Memo", "Debit", "Credit"]].copy()
        display["Date"] = display["Date"].astype(str).str[:10]
        for c in ["Debit", "Credit"]:
            display[c] = display[c].map(lambda v: f"${v:,.2f}" if v else "")
        st.dataframe(display, use_container_width=True, hide_index=True)


# ─────────────────────────────────────────────────────────────
# 3b · CLOSING-MONTH vs PRECEDING VARIANCE REPORTS
# ─────────────────────────────────────────────────────────────

# Balance-sheet keywords (assets, liabilities, equity)
_BS_KW = re.compile(
    r"(?i)(cash|bank|checking|savings|accounts?\s*receivable|a/?r|"
    r"inventory|prepaid|fixed\s*asset|equipment|furniture|"
    r"accumulated\s*depreciation|depreciation|"
    r"accounts?\s*payable|a/?p|accrued|loan|note|"
    r"credit\s*card|line\s*of\s*credit|mortgage|"
    r"equity|retained\s*earnings|owner|capital|"
    r"undeposited|other\s*current\s*asset|other\s*asset|"
    r"other\s*current\s*liabilit)"
)

# P&L keywords (revenue + expenses)
_PL_KW = re.compile(
    r"(?i)(income|revenue|sales|service|"
    r"cost\s*of\s*goods|cogs|cost\s*of\s*sales|"
    r"expense|rent|utilities|payroll|wage|salar|"
    r"insurance|supplies|office|travel|meal|"
    r"advertising|marketing|professional\s*fee|"
    r"repair|maintenance|tax|interest\s*expense|"
    r"depreciation\s*expense|amortization|"
    r"telephone|internet|shipping|freight|"
    r"commission|contractor|subcontract|"
    r"dues|subscription|license|training|"
    r"auto|vehicle|fuel|gas|"
    r"misc|other\s*income|other\s*expense|"
    r"gain|loss|discount)"
)

_MONTH_NAMES = ["January","February","March","April","May","June",
                "July","August","September","October","November","December"]


def _classify_account(acct_name: str) -> str:
    """Return 'BS', 'PL', or 'Unknown' based on account name keywords."""
    if _BS_KW.search(acct_name):
        return "BS"
    if _PL_KW.search(acct_name):
        return "PL"
    return "Unknown"


def _get_preceding_months(closing_month: int, start_month: int) -> list[int]:
    """Return the list of calendar months from start_month up to (but NOT including) closing_month.

    Wraps around December→January correctly.
    Example: start=7, closing=3 → [7,8,9,10,11,12,1,2]
    """
    months = []
    m = start_month
    while m != closing_month:
        months.append(m)
        m = (m % 12) + 1
    return months


def _build_balance_variance_table(balance_df: pd.DataFrame,
                                   closing_month: int, start_month: int,
                                   top_n: int = 20):
    """Build variance from an uploaded balance report (BS or P&L).

    balance_df has columns: Account, plus 'YYYY-MM' columns with balances.
    Returns (summary_df, preceding_labels, closing_label) or (None, None, None).
    """
    if balance_df is None or balance_df.empty:
        return None, None, None

    month_cols = sorted([c for c in balance_df.columns if c != "Account" and re.match(r"\d{4}-\d{2}", c)])
    if len(month_cols) < 2:
        return None, None, None

    # Map each column to its month number
    col_months = {}
    for c in month_cols:
        yr, mon = int(c.split("-")[0]), int(c.split("-")[1])
        col_months[c] = (yr, mon)

    preceding_month_nums = _get_preceding_months(closing_month, start_month)

    # Find closing and preceding columns
    closing_cols = [c for c in month_cols if col_months[c][1] == closing_month]
    preceding_cols = sorted([c for c in month_cols if col_months[c][1] in preceding_month_nums])

    if not closing_cols or not preceding_cols:
        return None, None, None

    close_col = closing_cols[-1]  # most recent
    closing_label = close_col

    # Build summary: Account | each preceding month balance | Preceding Avg | Closing | Variance | Var %
    records = []
    for _, row in balance_df.iterrows():
        acct = row["Account"]
        rec = {"Account": acct}
        prec_vals = []
        for pc in preceding_cols:
            val = row.get(pc, 0)
            rec[pc] = val
            prec_vals.append(val)
        prec_avg = np.mean(prec_vals) if prec_vals else 0
        close_val = row.get(close_col, 0)
        variance = close_val - prec_avg
        var_pct = (variance / abs(prec_avg)) if abs(prec_avg) > 0.01 else np.nan

        rec["Preceding Avg"] = prec_avg
        rec[f"Closing ({closing_label})"] = close_val
        rec["Variance ($)"] = variance
        rec["Variance (%)"] = var_pct
        rec["_abs_var"] = abs(variance)
        records.append(rec)

    summary = pd.DataFrame(records)
    summary = summary.sort_values("_abs_var", ascending=False).head(top_n)
    summary = summary.drop(columns=["_abs_var"])
    summary = summary.reset_index(drop=True)

    return summary, preceding_cols, closing_label


def _build_closing_variance_table_from_gl(gl: pd.DataFrame, account_type: str,
                                           closing_month: int, start_month: int,
                                           top_n: int = 20):
    """Fallback: build variance from GL transactions when no balance file uploaded."""
    if "YearMonth" not in gl.columns or gl["YearMonth"].isna().all():
        return None, None, None

    gl = gl.copy()
    if "AcctClass" not in gl.columns:
        gl["AcctClass"] = gl["Account"].apply(_classify_account)

    subset = gl[gl["AcctClass"] == account_type]
    if subset.empty:
        return None, None, None

    all_periods = sorted(subset["YearMonth"].dropna().unique())
    if len(all_periods) < 2:
        return None, None, None

    preceding_month_nums = _get_preceding_months(closing_month, start_month)
    if not preceding_month_nums:
        return None, None, None

    closing_periods = [p for p in all_periods if p.month == closing_month]
    preceding_periods = [p for p in all_periods if p.month in preceding_month_nums]
    if not closing_periods or not preceding_periods:
        return None, None, None

    close_p = closing_periods[-1]
    closing_label = str(close_p)
    preceding_periods = sorted(preceding_periods)
    preceding_labels = [str(p) for p in preceding_periods]

    all_relevant = preceding_periods + [close_p]
    monthly = (
        subset[subset["YearMonth"].isin(all_relevant)]
        .groupby(["Account", "YearMonth"])["Amount"]
        .sum()
        .unstack(fill_value=0)
    )

    prec_cols = [p for p in preceding_periods if p in monthly.columns]
    monthly["Preceding Avg"] = monthly[prec_cols].mean(axis=1) if prec_cols else 0
    monthly["Closing Month"] = monthly.get(close_p, 0)
    monthly["Variance"] = monthly["Closing Month"] - monthly["Preceding Avg"]
    monthly["Var_%"] = np.where(
        monthly["Preceding Avg"].abs() > 0.01,
        monthly["Variance"] / monthly["Preceding Avg"].abs(), np.nan)
    monthly["Abs_Var"] = monthly["Variance"].abs()
    monthly = monthly.sort_values("Abs_Var", ascending=False).head(top_n)

    records = []
    for acct in monthly.index:
        row = {"Account": acct}
        for p in preceding_periods:
            row[str(p)] = monthly.at[acct, p] if p in monthly.columns else 0
        row["Preceding Avg"] = monthly.at[acct, "Preceding Avg"]
        row[f"Closing ({closing_label})"] = monthly.at[acct, "Closing Month"]
        row["Variance ($)"] = monthly.at[acct, "Variance"]
        row["Variance (%)"] = monthly.at[acct, "Var_%"]
        records.append(row)

    return pd.DataFrame(records), preceding_labels, closing_label


def _render_variance_report(summary, prec_labels, close_label, closing_month, start_month, report_type, source_label):
    """Shared rendering for both BS and PL variance reports."""
    close_name = _MONTH_NAMES[closing_month - 1]
    start_name = _MONTH_NAMES[start_month - 1]
    prev_name = _MONTH_NAMES[((closing_month - 2) % 12)]
    type_label = "Balance Sheet" if report_type == "BS" else "Profit & Loss"

    section(f"Top 20 {type_label} — {close_name} Close vs {start_name}–{prev_name} Avg")

    if summary is None or summary.empty:
        st.warning(f"Could not build {type_label} variance analysis. "
                   f"Upload a {type_label} balance file in the sidebar, or ensure the GL covers the selected months.")
        return

    st.markdown(f"**Source:** {source_label}")
    st.markdown(f"**Closing month:** {close_label} &nbsp;|&nbsp; "
                f"**Preceding months:** {', '.join(str(p) for p in prec_labels)}")
    st.markdown("**Method:** Closing month balance compared to average of preceding months' balances")

    # Metric cards
    c1, c2, c3 = st.columns(3)
    top_acct = summary.iloc[0]["Account"] if len(summary) > 0 else "N/A"
    max_var = summary.iloc[0]["Variance ($)"] if len(summary) > 0 else 0
    avg_var = summary["Variance ($)"].abs().mean()
    with c1:
        st.markdown(metric_card("Largest Variance", f"${abs(max_var):,.2f}"), unsafe_allow_html=True)
    with c2:
        st.markdown(metric_card("Top Account", str(top_acct)[:35]), unsafe_allow_html=True)
    with c3:
        st.markdown(metric_card("Avg Abs Variance", f"${avg_var:,.2f}"), unsafe_allow_html=True)

    # Narrative
    section("Variance Highlights")
    for _, row in summary.head(5).iterrows():
        acct = row["Account"]
        var = row["Variance ($)"]
        pct = row.get("Variance (%)", np.nan)
        direction = "increased" if var > 0 else "decreased"
        pct_str = f" ({abs(pct):.1%})" if pd.notna(pct) else ""
        narrative(
            f"<strong>{acct}</strong> {direction} by <strong>${abs(var):,.2f}</strong>{pct_str} "
            f"in {close_label} compared to the preceding-months average of "
            f"${row['Preceding Avg']:,.2f}."
        )

    # Detail table
    section("Full Variance Detail")
    display = summary.copy()
    for c in display.columns:
        if c == "Account":
            continue
        if c == "Variance (%)":
            display[c] = display[c].map(lambda v: f"{v:.1%}" if pd.notna(v) else "—")
        else:
            display[c] = display[c].map(lambda v: f"${v:,.2f}" if isinstance(v, (int, float)) else v)
    st.dataframe(display, use_container_width=True, hide_index=True)


def report_bs_variance(gl: pd.DataFrame, closing_month: int, start_month: int, bs_balances=None):
    """Report 7 — Top 20 Balance Sheet: uploaded balances or GL fallback."""
    if bs_balances is not None and not bs_balances.empty:
        summary, prec_labels, close_label = _build_balance_variance_table(
            bs_balances, closing_month, start_month, 20)
        source = "Uploaded Balance Sheet report"
    else:
        summary, prec_labels, close_label = _build_closing_variance_table_from_gl(
            gl, "BS", closing_month, start_month, 20)
        source = "Derived from General Ledger transactions (upload a BS balance file for actual balances)"
    _render_variance_report(summary, prec_labels, close_label, closing_month, start_month, "BS", source)


def report_pl_variance(gl: pd.DataFrame, closing_month: int, start_month: int):
    """Report 8 — Top 20 P&L: derived from GL transactions."""
    summary, prec_labels, close_label = _build_closing_variance_table_from_gl(
        gl, "PL", closing_month, start_month, 20)
    _render_variance_report(summary, prec_labels, close_label, closing_month, start_month, "PL",
                            "Derived from General Ledger transactions")


# ─────────────────────────────────────────────────────────────
# 4 · IIF EXPORT ENGINE
# ─────────────────────────────────────────────────────────────

def generate_iif(gl: pd.DataFrame, period=None) -> str:
    """Generate a QuickBooks-importable IIF file from GL transactions."""
    lines = []
    lines.append("!TRNS\tTRNSID\tTRNSTYPE\tDATE\tACCNT\tNAME\tAMOUNT\tMEMO")
    lines.append("!SPL\tSPLID\tTRNSTYPE\tDATE\tACCNT\tNAME\tAMOUNT\tMEMO")
    lines.append("!ENDTRNS")

    if period is not None and "YearMonth" in gl.columns:
        data = gl[gl["YearMonth"] == period].copy()
    else:
        data = gl.copy()

    # Group by date + type to create journal entries
    if "Type" in data.columns:
        groups = data.groupby(["Date", "Type"])
    else:
        groups = data.groupby(["Date"])

    trns_id = 0
    for key, grp in groups:
        if len(grp) == 0:
            continue
        trns_id += 1
        first = grp.iloc[0]
        date_str = str(first["Date"])[:10] if pd.notna(first["Date"]) else ""
        trns_type = str(first.get("Type", "GENERAL JOURNAL")).strip()
        if trns_type in ("", "nan"):
            trns_type = "GENERAL JOURNAL"

        # First line = TRNS (the main entry)
        main_acct = str(first.get("Account", "")).replace("\t", " ")
        main_name = str(first.get("Name", "")).replace("\t", " ")
        if main_name == "nan":
            main_name = ""
        main_memo = str(first.get("Memo", "")).replace("\t", " ")
        if main_memo == "nan":
            main_memo = ""
        main_amount = first.get("Amount", 0)

        lines.append(f"TRNS\t{trns_id}\t{trns_type}\t{date_str}\t{main_acct}\t{main_name}\t{main_amount:.2f}\t{main_memo}")

        # Remaining lines = SPL
        for i in range(1, len(grp)):
            row = grp.iloc[i]
            acct = str(row.get("Account", "")).replace("\t", " ")
            name = str(row.get("Name", "")).replace("\t", " ")
            if name == "nan":
                name = ""
            memo = str(row.get("Memo", "")).replace("\t", " ")
            if memo == "nan":
                memo = ""
            amt = row.get("Amount", 0)
            lines.append(f"SPL\t{trns_id}\t{trns_type}\t{date_str}\t{acct}\t{name}\t{amt:.2f}\t{memo}")

        lines.append("ENDTRNS")

    return "\n".join(lines)


# ─────────────────────────────────────────────────────────────
# 5 · WORD DOCUMENT EXPORT ENGINE (python-docx)
# ─────────────────────────────────────────────────────────────

from docx import Document as DocxDocument
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def generate_docx_report(report_data: dict, output_path: str):
    """
    Generate an Apple-branded Word document from report data using python-docx.
    report_data keys:
      - title: str
      - subtitle: str
      - date: str
      - sections: list of {heading, paragraphs: list[str], table: {headers, rows}}
      - checks: list of {check, status, detail}  (for preflight)
    """
    doc = DocxDocument()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # ── Default font ──
    style = doc.styles["Normal"]
    style.font.name = "Helvetica"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x1D, 0x1D, 0x1F)

    # ── Header ──
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.clear()
    run1 = hp.add_run("Organization")
    run1.bold = True
    run1.font.size = Pt(8)
    run1.font.color.rgb = RGBColor(0x8E, 0x8E, 0x93)
    run2 = hp.add_run("  |  Month-End Close Agent")
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(0x8E, 0x8E, 0x93)

    # ── Footer ──
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Organization  |  Finance Department  |  Month-End Close Agent")
    fr.font.size = Pt(7)
    fr.font.color.rgb = RGBColor(0x8E, 0x8E, 0x93)

    # ── Title block ──
    doc.add_paragraph()  # spacer
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(report_data.get("title", "Report"))
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(0x1D, 0x1D, 0x1F)

    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run(report_data.get("subtitle", ""))
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(0x6E, 0x6E, 0x73)

    date_str = report_data.get("date", datetime.date.today().strftime("%d %b %Y"))
    date_p = doc.add_paragraph()
    date_run = date_p.add_run(f"Generated: {date_str}")
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(0x8E, 0x8E, 0x93)

    # Blue divider after title
    border_p = doc.add_paragraph()
    pPr = border_p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}>'
                     f'<w:bottom w:val="single" w:sz="12" w:space="4" w:color="0071E3"/>'
                     f'</w:pBdr>')
    pPr.append(pBdr)

    doc.add_paragraph()  # spacer

    # ── Sections ──
    for sec_data in report_data.get("sections", []):
        heading_text = sec_data.get("heading", "")
        if heading_text:
            h_p = doc.add_paragraph()
            h_run = h_p.add_run(heading_text)
            h_run.bold = True
            h_run.font.size = Pt(14)
            h_run.font.color.rgb = RGBColor(0x00, 0x71, 0xE3)
            # Blue underline on heading
            hPr = h_p._p.get_or_add_pPr()
            hBdr = parse_xml(f'<w:pBdr {nsdecls("w")}>'
                             f'<w:bottom w:val="single" w:sz="8" w:space="4" w:color="0071E3"/>'
                             f'</w:pBdr>')
            hPr.append(hBdr)

        for para_text in sec_data.get("paragraphs", []):
            p = doc.add_paragraph()
            r = p.add_run(str(para_text))
            r.font.size = Pt(10.5)

        # Table
        table_data = sec_data.get("table")
        if table_data and table_data.get("headers") and table_data.get("rows"):
            headers = table_data["headers"]
            rows = table_data["rows"]
            n_cols = len(headers)
            tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl.autofit = True

            # Style header row
            for i, h in enumerate(headers):
                cell = tbl.rows[0].cells[i]
                cell.text = ""
                p = cell.paragraphs[0]
                r = p.add_run(str(h))
                r.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                # Blue background
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="0071E3" w:val="clear"/>')
                cell._tc.get_or_add_tcPr().append(shading)

            # Data rows with alternating shading
            for ri, row in enumerate(rows):
                fill = "F5F5F7" if ri % 2 == 0 else "FFFFFF"
                for ci, cell_val in enumerate(row):
                    if ci >= n_cols:
                        break
                    cell = tbl.rows[ri + 1].cells[ci]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    r = p.add_run(str(cell_val) if cell_val is not None else "")
                    r.font.size = Pt(9)
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}" w:val="clear"/>')
                    cell._tc.get_or_add_tcPr().append(shading)

            # Light gray borders on all cells
            for row_obj in tbl.rows:
                for cell in row_obj.cells:
                    tcPr = cell._tc.get_or_add_tcPr()
                    borders = parse_xml(
                        f'<w:tcBorders {nsdecls("w")}>'
                        f'<w:top w:val="single" w:sz="2" w:space="0" w:color="D2D2D7"/>'
                        f'<w:bottom w:val="single" w:sz="2" w:space="0" w:color="D2D2D7"/>'
                        f'<w:left w:val="single" w:sz="2" w:space="0" w:color="D2D2D7"/>'
                        f'<w:right w:val="single" w:sz="2" w:space="0" w:color="D2D2D7"/>'
                        f'</w:tcBorders>'
                    )
                    tcPr.append(borders)

            doc.add_paragraph()  # spacer after table

    # ── Checks section (for preflight report) ──
    checks = report_data.get("checks", [])
    if checks:
        doc.add_paragraph()
        for ck in checks:
            icon = "\u2705" if ck["status"] == "PASS" else ("\u274C" if ck["status"] == "FAIL" else "\u26A0\uFE0F")
            ck_text = f'{icon} {ck["check"]} \u2014 {ck["detail"]}'
            p = doc.add_paragraph()
            r = p.add_run(ck_text)
            r.font.size = Pt(10)

    doc.save(output_path)


def build_flux_docx_data(gl: pd.DataFrame) -> dict:
    """Build report_data dict for the Flux Narrative report."""
    if "YearMonth" not in gl.columns:
        return None
    periods = sorted(gl["YearMonth"].dropna().unique())
    if len(periods) < 2:
        return None

    curr, prev = periods[-1], periods[-2]
    curr_df = gl[gl["YearMonth"] == curr]
    prev_df = gl[gl["YearMonth"] == prev]

    curr_totals = curr_df.groupby("Account")["Amount"].sum()
    prev_totals = prev_df.groupby("Account")["Amount"].sum()

    flux = pd.DataFrame({"Current": curr_totals, "Prior": prev_totals}).fillna(0)
    flux["Variance"] = flux["Current"] - flux["Prior"]
    flux["Var_%"] = np.where(flux["Prior"] != 0, (flux["Variance"] / flux["Prior"].abs()) * 100, np.nan)
    flux = flux.sort_values("Variance", key=abs, ascending=False)

    # Narrative paragraphs
    narratives = []
    for acct, row in flux.head(10).iterrows():
        direction = "increased" if row["Variance"] > 0 else "decreased"
        pct = f"{abs(row['Var_%']):.1f}%" if pd.notna(row["Var_%"]) else "N/A"
        acct_txns = curr_df[curr_df["Account"] == acct]
        vendors = [str(v) for v in acct_txns["Name"].dropna().unique()[:3] if str(v).strip() and str(v) != "nan"]
        memos = [str(m) for m in acct_txns["Memo"].dropna().unique()[:3] if str(m).strip() and str(m) != "nan"]
        vendor_str = ", ".join(vendors) if vendors else "various"
        memo_str = "; ".join(memos) if memos else "no memo detail"
        narratives.append(
            f"{acct} {direction} by ${abs(row['Variance']):,.2f} ({pct}), "
            f"from ${row['Prior']:,.2f} to ${row['Current']:,.2f}. "
            f"Key vendors: {vendor_str}. Memos: {memo_str}."
        )

    # Table
    table_rows = []
    for acct, row in flux.head(30).iterrows():
        table_rows.append([
            str(acct)[:40],
            f"${row['Current']:,.2f}",
            f"${row['Prior']:,.2f}",
            f"${row['Variance']:,.2f}",
            f"{row['Var_%']:.1f}%" if pd.notna(row["Var_%"]) else "N/A",
        ])

    return {
        "title": "Flux (Variance) Narrative Report",
        "subtitle": f"Organization  |  {prev} vs {curr}",
        "date": datetime.date.today().strftime("%d %b %Y"),
        "sections": [
            {"heading": "Variance Narrative", "paragraphs": narratives},
            {"heading": "Variance Detail",
             "paragraphs": [],
             "table": {
                 "headers": ["Account", "Current", "Prior", "Variance", "Var %"],
                 "rows": table_rows
             }},
        ]
    }


def build_vendor_gap_docx_data(gl: pd.DataFrame, closing_month: int = 3, start_month: int = 7) -> dict:
    """Build report_data dict for Missing Bill report with per-month detail."""
    if "Name" not in gl.columns or "YearMonth" not in gl.columns:
        return None
    vendors = gl[gl["Name"].str.strip() != "nan"].copy()
    if vendors.empty:
        return None
    all_periods = sorted(vendors["YearMonth"].dropna().unique())
    if len(all_periods) < 3:
        return None

    closing_periods = [p for p in all_periods if p.month == closing_month]
    prec_month_nums = _get_preceding_months(closing_month, start_month)
    preceding_periods = sorted([p for p in all_periods if p.month in prec_month_nums])
    close_p = closing_periods[-1] if closing_periods else all_periods[-1]
    analysis_periods = preceding_periods + [close_p]

    vendors_in_scope = vendors[vendors["YearMonth"].isin(analysis_periods)]
    history = vendors_in_scope[vendors_in_scope["YearMonth"] != close_p]
    curr_vendors = set(vendors_in_scope[vendors_in_scope["YearMonth"] == close_p]["Name"].unique())
    total_hist = len(preceding_periods)
    vendor_periods = history.groupby("Name")["YearMonth"].nunique()
    recurring = vendor_periods[vendor_periods >= max(2, total_hist * 0.5)]
    missing = [v for v in recurring.index if v not in curr_vendors and str(v).strip() and str(v) != "nan"]

    accrual_rows = []
    for v in missing:
        v_hist = history[history["Name"] == v]
        avg_amt = v_hist["Credit"].mean()
        last_date = v_hist["Date"].max()
        freq = v_hist["YearMonth"].nunique()
        typical_acct = v_hist["Account"].mode().iloc[0] if not v_hist["Account"].mode().empty else "Unknown"
        accrual_rows.append([
            str(v), f"{freq}/{total_hist} months",
            f"${avg_amt:,.2f}" if pd.notna(avg_amt) else "N/A",
            str(last_date.date()) if pd.notna(last_date) else "N/A",
            str(typical_acct)[:40],
            f"${avg_amt:,.2f}" if pd.notna(avg_amt) else "N/A",
        ])

    # Per-month transaction detail
    month_headers = ["Vendor"] + [str(p) for p in analysis_periods]
    month_rows = []
    for v in missing:
        v_all = vendors_in_scope[vendors_in_scope["Name"] == v]
        r = [str(v)]
        for p in analysis_periods:
            total = v_all[v_all["YearMonth"] == p]["Credit"].sum()
            r.append(f"${total:,.2f}" if total != 0 else "—")
        month_rows.append(r)

    return {
        "title": "Recurring Vendor Gap Analysis",
        "subtitle": f"Organization  |  Missing Bill Report  |  Closing: {close_p}",
        "date": datetime.date.today().strftime("%d %b %Y"),
        "sections": [
            {"heading": "Summary",
             "paragraphs": [
                 f"{len(recurring)} recurring vendors across {total_hist} preceding months.",
                 f"{len(missing)} vendor(s) missing from closing period ({close_p}) — may need accrual entries.",
             ]},
            {"heading": "Suggested Accruals", "paragraphs": [],
             "table": {
                 "headers": ["Vendor", "Frequency", "Avg Amount", "Last Seen", "Typical Account", "Suggested Accrual"],
                 "rows": accrual_rows}},
            {"heading": "Monthly Transaction Detail", "paragraphs": [],
             "table": {"headers": month_headers, "rows": month_rows}},
        ]
    }


def build_suspense_docx_data(gl: pd.DataFrame, coa: dict | None) -> dict:
    """Build report_data for Suspense & Misc Reclass."""
    patterns = r"(?i)(suspense|misc|other|unclass|uncategoriz|clearing|unknown|unallocat)"
    mask = gl["Account"].str.contains(patterns, na=False) | gl["Memo"].str.contains(patterns, na=False)
    if "Split" in gl.columns:
        mask = mask | gl["Split"].str.contains(patterns, na=False)
    flagged = gl[mask].copy()

    active_accts = sorted(coa.get("active_accounts", set())) if coa else []
    rows = []
    for _, txn in flagged.iterrows():
        memo = str(txn.get("Memo", "")).lower()
        name = str(txn.get("Name", "")).lower()
        search_text = memo + " " + name
        best_match, best_score = "", 0
        for acct in active_accts:
            keywords = [k for k in re.split(r"[:\s\-&/]+", acct.lower()) if len(k) > 2]
            score = sum(1 for k in keywords if k in search_text)
            if score > best_score:
                best_score = score
                best_match = acct
        rows.append([
            str(txn.get("Date", ""))[:10],
            str(txn.get("Account", ""))[:35],
            str(txn.get("Name", ""))[:25],
            str(txn.get("Memo", ""))[:40],
            f"${txn['Amount']:,.2f}",
            best_match[:40] if best_score >= 1 else "Manual review",
            "High" if best_score >= 3 else ("Medium" if best_score >= 1 else "Low"),
        ])

    return {
        "title": "Suspense & Misc Resolution Worksheet",
        "subtitle": "Organization  |  Reclassification Recommendations",
        "date": datetime.date.today().strftime("%d %b %Y"),
        "sections": [
            {"heading": "Summary",
             "paragraphs": [
                 f"{len(flagged)} transaction(s) flagged in suspense, misc, clearing, or unclassified accounts.",
                 f"Total amount at risk: ${flagged['Amount'].abs().sum():,.2f}.",
             ]},
            {"heading": "Resolution Worksheet",
             "paragraphs": [],
             "table": {
                 "headers": ["Date", "Current Account", "Name", "Memo", "Amount", "Suggested Reclass", "Confidence"],
                 "rows": rows[:50]
             }},
        ]
    }


def build_materiality_docx_data(gl: pd.DataFrame, threshold: float) -> dict:
    """Build report_data for Materiality & Risk report."""
    large = gl[gl["Amount"].abs() >= threshold].copy()
    rows = []
    for _, txn in large.iterrows():
        risk = "Low"
        reasons = []
        acct = str(txn.get("Account", "")).lower()
        memo = str(txn.get("Memo", "")).lower()
        if re.search(r"(suspense|misc|other|clearing|unknown)", acct):
            risk = "High"
            reasons.append("suspense/misc account")
        if memo in ("", "nan", "none"):
            risk = "High" if risk == "High" else "Medium"
            reasons.append("no memo")
        if abs(txn["Amount"]) >= threshold * 5:
            risk = "High"
            reasons.append("exceeds 5x threshold")
        rows.append([
            str(txn.get("Date", ""))[:10],
            str(txn.get("Account", ""))[:35],
            str(txn.get("Name", ""))[:25],
            f"${txn['Amount']:,.2f}",
            risk,
            "; ".join(reasons) if reasons else "Material amount",
        ])

    return {
        "title": "Materiality & Risk Threshold Report",
        "subtitle": f"Organization  |  Threshold: ${threshold:,.0f}",
        "date": datetime.date.today().strftime("%d %b %Y"),
        "sections": [
            {"heading": "Summary",
             "paragraphs": [
                 f"Materiality threshold set at ${threshold:,.0f}.",
                 f"{len(rows)} transaction(s) exceed the threshold and have been flagged for review.",
                 f"High risk items: {sum(1 for r in rows if r[4] == 'High')}. Medium: {sum(1 for r in rows if r[4] == 'Medium')}.",
             ]},
            {"heading": "Flagged Transactions",
             "paragraphs": [],
             "table": {
                 "headers": ["Date", "Account", "Name", "Amount", "Risk", "Reason"],
                 "rows": rows[:50]
             }},
        ]
    }


def build_preflight_docx_data(gl: pd.DataFrame, coa: dict | None) -> dict:
    """Build report_data for IIF Pre-Flight."""
    checks = []
    if "YearMonth" in gl.columns:
        for period in sorted(gl["YearMonth"].dropna().unique()):
            p_df = gl[gl["YearMonth"] == period]
            total_dr, total_cr = p_df["Debit"].sum(), p_df["Credit"].sum()
            diff = abs(total_dr - total_cr)
            checks.append({
                "check": f"Debits = Credits ({period})",
                "status": "PASS" if diff < 0.01 else "FAIL",
                "detail": f"DR: ${total_dr:,.2f} | CR: ${total_cr:,.2f} | Diff: ${diff:,.2f}",
            })

    blank_accts = gl["Account"].isna().sum() + (gl["Account"] == "").sum() + (gl["Account"] == "nan").sum()
    checks.append({
        "check": "No blank account codes",
        "status": "PASS" if blank_accts == 0 else "WARN",
        "detail": f"{blank_accts} blank" if blank_accts else "All accounts populated",
    })

    if "Date" in gl.columns:
        future = gl[gl["Date"] > pd.Timestamp.now()].shape[0]
        checks.append({
            "check": "No future-dated transactions",
            "status": "PASS" if future == 0 else "WARN",
            "detail": f"{future} future-dated" if future else "All dates current",
        })

    return {
        "title": "IIF Import Pre-Flight Validation",
        "subtitle": "Organization  |  Technical Validation",
        "date": datetime.date.today().strftime("%d %b %Y"),
        "sections": [
            {"heading": "Validation Results",
             "paragraphs": [f"{sum(1 for c in checks if c['status']=='PASS')}/{len(checks)} checks passed."]}
        ],
        "checks": checks,
    }


def _build_variance_docx_data(gl: pd.DataFrame, account_type: str, title: str,
                               closing_month: int, start_month: int,
                               balance_df=None) -> dict:
    """Build report_data for closing-vs-preceding variance report (BS or PL)."""
    if balance_df is not None and not balance_df.empty:
        summary, prec_labels, close_label = _build_balance_variance_table(
            balance_df, closing_month, start_month, 20)
    else:
        summary, prec_labels, close_label = _build_closing_variance_table_from_gl(
            gl, account_type, closing_month, start_month, 20)
    if summary is None or summary.empty:
        return None

    close_name = _MONTH_NAMES[closing_month - 1]
    start_name = _MONTH_NAMES[start_month - 1]

    narratives = []
    for _, row in summary.head(5).iterrows():
        acct = row["Account"]
        var = row["Variance ($)"]
        pct = row.get("Variance (%)", np.nan)
        direction = "increased" if var > 0 else "decreased"
        pct_str = f" ({abs(pct):.1%})" if pd.notna(pct) else ""
        narratives.append(
            f"{acct} {direction} by ${abs(var):,.2f}{pct_str} in {close_label} "
            f"compared to the preceding-months average of ${row['Preceding Avg']:,.2f}."
        )

    # Table headers: Account, each preceding month, Preceding Avg, Closing, Variance, Var %
    headers = ["Account"] + prec_labels + ["Preceding Avg", f"Closing ({close_label})", "Variance ($)", "Variance (%)"]
    table_rows = []
    for _, row in summary.iterrows():
        r = [str(row["Account"])[:40]]
        for pl in prec_labels:
            v = row.get(pl, 0)
            r.append(f"${v:,.2f}" if isinstance(v, (int, float)) else str(v))
        r.append(f"${row['Preceding Avg']:,.2f}")
        r.append(f"${row[f'Closing ({close_label})']:,.2f}")
        r.append(f"${row['Variance ($)']:,.2f}")
        pct = row.get("Variance (%)", np.nan)
        r.append(f"{pct:.1%}" if pd.notna(pct) else "N/A")
        table_rows.append(r)

    return {
        "title": title,
        "subtitle": f"Organization  |  {close_name} Close vs {start_name}–{_MONTH_NAMES[((closing_month - 2) % 12)]} Avg",
        "date": datetime.date.today().strftime("%d %b %Y"),
        "sections": [
            {"heading": "Summary",
             "paragraphs": [
                 f"Top 20 {account_type} accounts ranked by largest variance.",
                 f"Closing month {close_label} compared to average of {len(prec_labels)} preceding months.",
             ]},
            {"heading": "Variance Highlights", "paragraphs": narratives},
            {"heading": "Full Variance Detail",
             "paragraphs": [],
             "table": {"headers": headers, "rows": table_rows}},
        ],
    }


def build_bs_variance_docx_data(gl, closing_month, start_month, bs_balances=None) -> dict:
    return _build_variance_docx_data(gl, "BS", "Top 20 Balance Sheet — Largest Variances",
                                      closing_month, start_month, bs_balances)


def build_pl_variance_docx_data(gl, closing_month, start_month) -> dict:
    return _build_variance_docx_data(gl, "PL", "Top 20 Profit & Loss — Largest Variances",
                                      closing_month, start_month, None)


def export_all_reports_docx(gl, coa, threshold, pdf_texts, closing_month=3, start_month=7,
                            bs_balances=None):
    """Generate all 8 reports as Word documents and return as a zip buffer."""
    import zipfile

    reports = {
        "01_Flux_Narrative": build_flux_docx_data(gl),
        "02_Missing_Bill_Analysis": build_vendor_gap_docx_data(gl, closing_month, start_month),
        "03_Suspense_Resolution": build_suspense_docx_data(gl, coa),
        "04_Materiality_Risk": build_materiality_docx_data(gl, threshold),
        "05_IIF_PreFlight": build_preflight_docx_data(gl, coa),
        "07_BS_Top20_Variance": build_bs_variance_docx_data(gl, closing_month, start_month, bs_balances),
        "08_PL_Top20_Variance": build_pl_variance_docx_data(gl, closing_month, start_month),
    }

    # Report 6 - Reconciliation summary (simpler, text-based)
    bank_kw = r"(?i)(bank|cash|checking|savings|operating|deposit)"
    bank_gl = gl[gl["Account"].str.contains(bank_kw, na=False)]
    tax_kw = r"(?i)(sales\s*tax|tax|shipping|freight|delivery|handling)"
    tax_txns = gl[gl["Memo"].str.contains(tax_kw, na=False) | gl["Account"].str.contains(tax_kw, na=False)]

    recon_rows = []
    for _, txn in tax_txns.head(30).iterrows():
        recon_rows.append([
            str(txn.get("Date", ""))[:10],
            str(txn.get("Account", ""))[:35],
            str(txn.get("Memo", ""))[:40],
            f"${txn.get('Debit', 0):,.2f}" if txn.get("Debit", 0) else "",
            f"${txn.get('Credit', 0):,.2f}" if txn.get("Credit", 0) else "",
        ])

    reports["06_Reconciliation_Summary"] = {
        "title": "Multi-Source Reconciliation Summary",
        "subtitle": "Organization  |  Three-Way Reconciliation",
        "date": datetime.date.today().strftime("%d %b %Y"),
        "sections": [
            {"heading": "Overview",
             "paragraphs": [
                 f"GL bank transactions identified: {len(bank_gl):,}.",
                 f"PDF documents uploaded for reconciliation: {len(pdf_texts)}.",
                 f"Sales tax and shipping entries found: {len(tax_txns)}.",
             ]},
            {"heading": "Tax & Shipping Line Items",
             "paragraphs": [],
             "table": {
                 "headers": ["Date", "Account", "Memo", "Debit", "Credit"],
                 "rows": recon_rows
             }},
        ]
    }

    with tempfile.TemporaryDirectory() as tmpdir:
        paths = {}
        for name, data in reports.items():
            if data is None:
                continue
            path = os.path.join(tmpdir, f"{name}.docx")
            try:
                generate_docx_report(data, path)
                paths[name] = path
            except Exception as e:
                st.warning(f"Could not generate {name}: {e}")

        # Bundle into zip
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
            for name, path in paths.items():
                zf.write(path, f"Month-End Reports/{name}.docx")

        zip_buffer.seek(0)
        return zip_buffer, list(paths.keys())


# ─────────────────────────────────────────────────────────────
# 5c · EXCEL EXPORT ENGINE
# ─────────────────────────────────────────────────────────────

def export_all_reports_xlsx(gl: pd.DataFrame, coa, threshold: float, pdf_texts: list,
                            closing_month: int = 3, start_month: int = 7,
                            bs_balances=None) -> io.BytesIO:
    """Generate all 8 reports as sheets in a single styled Excel workbook."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    # Remove the default sheet — we'll create named ones
    wb.remove(wb.active)

    # ── Style constants ──
    hdr_font = Font(name="Arial", bold=True, size=10, color="FFFFFF")
    hdr_fill = PatternFill("solid", fgColor="0071E3")
    hdr_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    data_font = Font(name="Arial", size=10)
    curr_fmt = '#,##0.00;(#,##0.00);"-"'
    pct_fmt = '0.0%;(0.0%);"-"'
    thin_border = Border(
        left=Side(style="thin", color="D2D2D7"),
        right=Side(style="thin", color="D2D2D7"),
        top=Side(style="thin", color="D2D2D7"),
        bottom=Side(style="thin", color="D2D2D7"),
    )
    alt_fill = PatternFill("solid", fgColor="F5F5F7")
    title_font = Font(name="Arial", bold=True, size=14, color="0071E3")
    subtitle_font = Font(name="Arial", size=10, color="6E6E73")

    def _add_title_block(ws, title, subtitle=""):
        ws.append([title])
        ws.cell(row=1, column=1).font = title_font
        ws.append([subtitle])
        ws.cell(row=2, column=1).font = subtitle_font
        ws.append([f"Generated {datetime.date.today().strftime('%d %b %Y')}"])
        ws.cell(row=3, column=1).font = subtitle_font
        ws.append([])  # blank row

    def _write_table(ws, headers, rows, start_row=None):
        if start_row is None:
            start_row = ws.max_row + 1
        # Headers
        for ci, h in enumerate(headers, 1):
            cell = ws.cell(row=start_row, column=ci, value=h)
            cell.font = hdr_font
            cell.fill = hdr_fill
            cell.alignment = hdr_align
            cell.border = thin_border
        # Data rows
        for ri, row_data in enumerate(rows):
            excel_row = start_row + 1 + ri
            for ci, val in enumerate(row_data, 1):
                cell = ws.cell(row=excel_row, column=ci, value=val)
                cell.font = data_font
                cell.border = thin_border
                cell.alignment = Alignment(horizontal="right" if ci > 1 else "left",
                                           vertical="center")
                if ri % 2 == 1:
                    cell.fill = alt_fill
        # Auto-width
        for ci in range(1, len(headers) + 1):
            max_len = len(str(headers[ci - 1]))
            for ri in range(len(rows)):
                cell_val = str(rows[ri][ci - 1]) if ci - 1 < len(rows[ri]) else ""
                max_len = max(max_len, len(cell_val))
            ws.column_dimensions[get_column_letter(ci)].width = min(max_len + 3, 40)
        return start_row + 1 + len(rows)

    def _add_section_heading(ws, text):
        row = ws.max_row + 2
        cell = ws.cell(row=row, column=1, value=text)
        cell.font = Font(name="Arial", bold=True, size=11, color="0071E3")
        return row

    # ────────────────────────────────────────
    # Sheet 1: Flux Narrative
    # ────────────────────────────────────────
    if "YearMonth" in gl.columns and not gl["YearMonth"].isna().all():
        periods = sorted(gl["YearMonth"].dropna().unique())
        if len(periods) >= 2:
            ws = wb.create_sheet("1-Flux Narrative")
            curr, prev = periods[-1], periods[-2]
            _add_title_block(ws, "Flux (Variance) Narrative", f"{prev} vs {curr}")

            curr_df = gl[gl["YearMonth"] == curr]
            prev_df = gl[gl["YearMonth"] == prev]
            curr_totals = curr_df.groupby("Account")["Amount"].sum()
            prev_totals = prev_df.groupby("Account")["Amount"].sum()
            flux = pd.DataFrame({"Current": curr_totals, "Prior": prev_totals}).fillna(0)
            flux["Variance"] = flux["Current"] - flux["Prior"]
            flux["Var_%"] = np.where(flux["Prior"] != 0, flux["Variance"] / flux["Prior"].abs(), np.nan)
            flux = flux.sort_values("Variance", key=abs, ascending=False)

            headers = ["Account", "Current ($)", "Prior ($)", "Variance ($)", "Var %"]
            rows = []
            for acct, r in flux.iterrows():
                rows.append([
                    str(acct),
                    round(r["Current"], 2),
                    round(r["Prior"], 2),
                    round(r["Variance"], 2),
                    round(r["Var_%"], 4) if pd.notna(r["Var_%"]) else None,
                ])
            _write_table(ws, headers, rows)
            # Format currency and percent columns
            for ri in range(len(rows)):
                for ci in [2, 3, 4]:
                    ws.cell(row=5 + ri, column=ci).number_format = curr_fmt
                ws.cell(row=5 + ri, column=5).number_format = pct_fmt

    # ────────────────────────────────────────
    # Sheet 2: Missing Bills
    # ────────────────────────────────────────
    if "Name" in gl.columns and "YearMonth" in gl.columns:
        vendors = gl[gl["Name"].str.strip() != "nan"].copy()
        if not vendors.empty:
            all_periods = sorted(vendors["YearMonth"].dropna().unique())
            closing_periods = [p for p in all_periods if p.month == closing_month]
            prec_month_nums = _get_preceding_months(closing_month, start_month)
            preceding_periods = sorted([p for p in all_periods if p.month in prec_month_nums])
            close_p = closing_periods[-1] if closing_periods else all_periods[-1]
            analysis_periods = preceding_periods + [close_p]

            vendors_in_scope = vendors[vendors["YearMonth"].isin(analysis_periods)]
            history = vendors_in_scope[vendors_in_scope["YearMonth"] != close_p]
            curr_vendors = set(vendors_in_scope[vendors_in_scope["YearMonth"] == close_p]["Name"].unique())
            total_hist = len(preceding_periods)

            if total_hist >= 2:
                ws = wb.create_sheet("2-Missing Bills")
                _add_title_block(ws, "Recurring Vendor Gap Analysis", f"Closing: {close_p}")
                vendor_periods = history.groupby("Name")["YearMonth"].nunique()
                recurring = vendor_periods[vendor_periods >= max(2, total_hist * 0.5)]
                missing = [v for v in recurring.index if v not in curr_vendors and str(v).strip() and str(v) != "nan"]

                # Accrual summary table
                headers = ["Vendor", "Frequency", "Avg Amount ($)", "Last Seen", "Typical Account", "Suggested Accrual ($)"]
                rows = []
                for v in missing:
                    v_hist = history[history["Name"] == v]
                    avg_amt = v_hist["Credit"].mean()
                    last_date = v_hist["Date"].max()
                    freq = v_hist["YearMonth"].nunique()
                    typical_acct = v_hist["Account"].mode().iloc[0] if not v_hist["Account"].mode().empty else "Unknown"
                    rows.append([
                        str(v), f"{freq}/{total_hist}",
                        round(avg_amt, 2) if pd.notna(avg_amt) else 0,
                        str(last_date.date()) if pd.notna(last_date) else "",
                        str(typical_acct),
                        round(avg_amt, 2) if pd.notna(avg_amt) else 0,
                    ])
                next_r = _write_table(ws, headers, rows)
                for ri in range(len(rows)):
                    for ci in [3, 6]:
                        ws.cell(row=5 + ri, column=ci).number_format = curr_fmt

                # Per-month transaction detail
                if missing:
                    _add_section_heading(ws, "Monthly Transaction Detail")
                    month_headers = ["Vendor"] + [str(p) for p in analysis_periods]
                    month_rows = []
                    for v in missing:
                        v_all = vendors_in_scope[vendors_in_scope["Name"] == v]
                        r = [str(v)]
                        for p in analysis_periods:
                            total = v_all[v_all["YearMonth"] == p]["Credit"].sum()
                            r.append(round(total, 2) if total != 0 else 0)
                        month_rows.append(r)
                    mr_start = _write_table(ws, month_headers, month_rows)
                    for ri in range(len(month_rows)):
                        for ci in range(2, len(month_headers) + 1):
                            ws.cell(row=mr_start - len(month_rows) + ri, column=ci).number_format = curr_fmt

    # ────────────────────────────────────────
    # Sheet 3: Suspense Reclass
    # ────────────────────────────────────────
    suspense_kw = r"(?i)(suspense|clearing|misc|miscellaneous|unclassified|uncategorized|ask\s*my\s*account|other)"
    susp_txns = gl[gl["Account"].str.contains(suspense_kw, na=False)]
    if not susp_txns.empty:
        ws = wb.create_sheet("3-Suspense Reclass")
        _add_title_block(ws, "Suspense & Misc Resolution Worksheet")
        headers = ["Date", "Account", "Name", "Memo", "Debit ($)", "Credit ($)", "Suggested Reclass"]
        rows = []
        for _, txn in susp_txns.head(100).iterrows():
            rows.append([
                str(txn.get("Date", ""))[:10],
                str(txn.get("Account", "")),
                str(txn.get("Name", "")),
                str(txn.get("Memo", "")),
                round(txn.get("Debit", 0), 2),
                round(txn.get("Credit", 0), 2),
                "",  # blank for user to fill
            ])
        _write_table(ws, headers, rows)
        for ri in range(len(rows)):
            for ci in [5, 6]:
                ws.cell(row=5 + ri, column=ci).number_format = curr_fmt

    # ────────────────────────────────────────
    # Sheet 4: Materiality & Risk
    # ────────────────────────────────────────
    ws = wb.create_sheet("4-Materiality Risk")
    _add_title_block(ws, "Materiality & Risk Threshold", f"Threshold: ${threshold:,.0f}")
    material = gl[gl["Amount"].abs() >= threshold].sort_values("Amount", key=abs, ascending=False)
    headers = ["Date", "Account", "Name", "Amount ($)", "Risk Level", "Reason"]
    rows = []
    for _, txn in material.head(100).iterrows():
        risk = "Medium"
        reason = "Material amount"
        if abs(txn["Amount"]) >= threshold * 5:
            risk = "High"
            reason = "Exceeds 5x threshold"
        elif abs(txn["Amount"]) >= threshold * 2:
            risk = "Medium"
            reason = "Exceeds 2x threshold"
        rows.append([
            str(txn.get("Date", ""))[:10],
            str(txn.get("Account", "")),
            str(txn.get("Name", "")),
            round(txn["Amount"], 2),
            risk,
            reason,
        ])
    _write_table(ws, headers, rows)
    for ri in range(len(rows)):
        ws.cell(row=5 + ri, column=4).number_format = curr_fmt

    # ────────────────────────────────────────
    # Sheet 5: IIF Pre-Flight
    # ────────────────────────────────────────
    ws = wb.create_sheet("5-IIF PreFlight")
    _add_title_block(ws, "IIF Import Pre-Flight Validation")
    checks = []
    if "YearMonth" in gl.columns:
        for period in sorted(gl["YearMonth"].dropna().unique()):
            p_df = gl[gl["YearMonth"] == period]
            total_dr, total_cr = p_df["Debit"].sum(), p_df["Credit"].sum()
            diff = abs(total_dr - total_cr)
            checks.append([
                f"Debits = Credits ({period})",
                "PASS" if diff < 0.01 else "FAIL",
                f"DR: ${total_dr:,.2f} | CR: ${total_cr:,.2f} | Diff: ${diff:,.2f}",
            ])
    blank_accts = gl["Account"].isna().sum() + (gl["Account"] == "").sum() + (gl["Account"] == "nan").sum()
    checks.append(["No Blank Accounts", "PASS" if blank_accts == 0 else "FAIL", f"{blank_accts} blank account(s)"])
    headers = ["Check", "Status", "Detail"]
    _write_table(ws, headers, checks)
    # Color the status cells
    pass_fill = PatternFill("solid", fgColor="D4EDDA")
    fail_fill = PatternFill("solid", fgColor="F8D7DA")
    for ri in range(len(checks)):
        cell = ws.cell(row=5 + ri, column=2)
        cell.fill = pass_fill if cell.value == "PASS" else fail_fill

    # ────────────────────────────────────────
    # Sheet 6: Reconciliation
    # ────────────────────────────────────────
    ws = wb.create_sheet("6-Reconciliation")
    _add_title_block(ws, "Multi-Source Reconciliation Summary")
    bank_kw = r"(?i)(bank|cash|checking|savings|operating|deposit)"
    bank_gl = gl[gl["Account"].str.contains(bank_kw, na=False)]
    tax_kw = r"(?i)(sales\s*tax|tax|shipping|freight|delivery|handling)"
    tax_txns = gl[gl["Memo"].str.contains(tax_kw, na=False) | gl["Account"].str.contains(tax_kw, na=False)]

    _add_section_heading(ws, "Overview")
    r = ws.max_row + 1
    ws.cell(row=r, column=1, value=f"GL bank transactions: {len(bank_gl):,}").font = data_font
    ws.cell(row=r + 1, column=1, value=f"PDFs uploaded: {len(pdf_texts)}").font = data_font
    ws.cell(row=r + 2, column=1, value=f"Tax/shipping entries: {len(tax_txns)}").font = data_font

    _add_section_heading(ws, "Tax & Shipping Line Items")
    headers = ["Date", "Account", "Memo", "Debit ($)", "Credit ($)"]
    rows = []
    for _, txn in tax_txns.head(50).iterrows():
        rows.append([
            str(txn.get("Date", ""))[:10],
            str(txn.get("Account", "")),
            str(txn.get("Memo", "")),
            round(txn.get("Debit", 0), 2),
            round(txn.get("Credit", 0), 2),
        ])
    if rows:
        sr = _write_table(ws, headers, rows)
        for ri in range(len(rows)):
            for ci in [4, 5]:
                ws.cell(row=ws.max_row - len(rows) + ri + 1, column=ci).number_format = curr_fmt

    # ────────────────────────────────────────
    # Sheet 7: BS Top 20 Variance (Closing vs Preceding Avg)
    # ────────────────────────────────────────
    if bs_balances is not None and not bs_balances.empty:
        summary_bs, prec_bs, close_bs = _build_balance_variance_table(bs_balances, closing_month, start_month, 20)
    else:
        summary_bs, prec_bs, close_bs = _build_closing_variance_table_from_gl(gl, "BS", closing_month, start_month, 20)
    if summary_bs is not None and not summary_bs.empty:
        close_name = _MONTH_NAMES[closing_month - 1]
        start_name = _MONTH_NAMES[start_month - 1]
        ws = wb.create_sheet("7-BS Top20 Variance")
        _add_title_block(ws, "Top 20 Balance Sheet — Largest Variances",
                         f"{close_name} Close vs {start_name}–{_MONTH_NAMES[((closing_month - 2) % 12)]} Avg")
        headers = list(summary_bs.columns)
        rows = []
        for _, row in summary_bs.iterrows():
            r = []
            for h in headers:
                v = row[h]
                if h == "Account":
                    r.append(str(v)[:40])
                elif h == "Variance (%)":
                    r.append(round(v, 4) if pd.notna(v) else None)
                else:
                    r.append(round(v, 2) if isinstance(v, (int, float)) else v)
            rows.append(r)
        _write_table(ws, headers, rows)
        for ri in range(len(rows)):
            for ci in range(2, len(headers) + 1):
                if headers[ci - 1] == "Variance (%)":
                    ws.cell(row=5 + ri, column=ci).number_format = pct_fmt
                else:
                    ws.cell(row=5 + ri, column=ci).number_format = curr_fmt

    # ────────────────────────────────────────
    # Sheet 8: P&L Top 20 Variance (Closing vs Preceding Avg)
    # ────────────────────────────────────────
    summary_pl, prec_pl, close_pl = _build_closing_variance_table_from_gl(gl, "PL", closing_month, start_month, 20)
    if summary_pl is not None and not summary_pl.empty:
        close_name = _MONTH_NAMES[closing_month - 1]
        start_name = _MONTH_NAMES[start_month - 1]
        ws = wb.create_sheet("8-PL Top20 Variance")
        _add_title_block(ws, "Top 20 Profit & Loss — Largest Variances",
                         f"{close_name} Close vs {start_name}–{_MONTH_NAMES[((closing_month - 2) % 12)]} Avg")
        headers = list(summary_pl.columns)
        rows = []
        for _, row in summary_pl.iterrows():
            r = []
            for h in headers:
                v = row[h]
                if h == "Account":
                    r.append(str(v)[:40])
                elif h == "Variance (%)":
                    r.append(round(v, 4) if pd.notna(v) else None)
                else:
                    r.append(round(v, 2) if isinstance(v, (int, float)) else v)
            rows.append(r)
        _write_table(ws, headers, rows)
        for ri in range(len(rows)):
            for ci in range(2, len(headers) + 1):
                if headers[ci - 1] == "Variance (%)":
                    ws.cell(row=5 + ri, column=ci).number_format = pct_fmt
                else:
                    ws.cell(row=5 + ri, column=ci).number_format = curr_fmt

    # ── Write to buffer ──
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


# ─────────────────────────────────────────────────────────────
# 6 · MAIN APPLICATION
# ─────────────────────────────────────────────────────────────

def main():
    # Header
    st.markdown("""
    <div class="apple-header">
        <h1>Month-End Close Agent</h1>
        <p>Organization &nbsp;|&nbsp; Finance Department &nbsp;|&nbsp; Diagnostic Reporting Suite</p>
    </div>
    """, unsafe_allow_html=True)

    # ── Sidebar ──
    with st.sidebar:
        st.markdown("### Data Inputs")
        st.markdown("Upload your accounting source files to generate diagnostic reports.")

        gl_file = st.file_uploader(
            "General Ledger (GL)",
            type=["csv", "xlsx", "xls"],
            help="Transaction history export from QuickBooks (CSV or Excel).",
        )
        coa_file = st.file_uploader(
            "Chart of Accounts & Classes",
            type=["iif", "csv"],
            help="IIF export or CSV with account names and active classes.",
        )
        pdf_files = st.file_uploader(
            "Invoices / Bank Statements (optional)",
            type=["pdf"],
            accept_multiple_files=True,
            help="Upload PDFs for multi-source reconciliation.",
        )

        st.markdown("---")
        st.markdown("### Balance Uploads")
        st.caption("Monthly balance reports for variance analysis (Reports 7 & 8).")
        bs_balance_file = st.file_uploader(
            "Balance Sheet Balances",
            type=["csv", "xlsx", "xls"],
            help="Monthly balance sheet report — accounts as rows, months as columns. "
                 "QuickBooks: Reports → Company & Financial → Balance Sheet Standard → Monthly.",
        )

        st.markdown("---")
        st.markdown("### Period Settings")

        _month_names = ["January","February","March","April","May","June",
                        "July","August","September","October","November","December"]
        closing_month = st.selectbox(
            "Closing Month",
            options=list(range(1, 13)),
            format_func=lambda m: _month_names[m - 1],
            index=2,  # default March
            help="The month you are closing. Variance analysis compares this month against the preceding months.",
        )
        start_month = st.selectbox(
            "Variance Start Month",
            options=list(range(1, 13)),
            format_func=lambda m: _month_names[m - 1],
            index=6,  # default July
            help="First month of the preceding period used for variance comparison (e.g. July for a fiscal year starting July 1).",
        )

        st.markdown("---")
        st.markdown("### Controls")
        threshold = st.slider(
            "Materiality Threshold ($)",
            min_value=100,
            max_value=50_000,
            value=1_000,
            step=100,
            help="Transactions exceeding this amount are flagged for audit risk review.",
        )

        st.markdown("---")
        st.markdown("### Export")
        export_iif = st.checkbox("Enable IIF Export", value=False)
        export_docx = st.checkbox("Export Reports as Word (.docx)", value=False,
                                  help="Generate Apple-branded Word documents for all 8 reports.")
        export_xlsx = st.checkbox("Export Reports as Excel (.xlsx)", value=False,
                                  help="Generate a multi-sheet Excel workbook with all 8 reports.")

    # ── Parse data ──
    gl = None
    coa = None
    pdf_texts = []
    bs_balances = None

    if gl_file:
        with st.spinner("Parsing General Ledger..."):
            gl = parse_gl(gl_file)
        st.sidebar.success(f"GL loaded — {len(gl):,} transactions")

    if coa_file:
        with st.spinner("Parsing Chart of Accounts..."):
            if coa_file.name.lower().endswith(".iif"):
                coa = parse_iif(coa_file)
            else:
                coa = parse_coa_csv(coa_file)
        n_accts = len(coa.get("active_accounts", []))
        n_cls = len(coa.get("active_classes", []))
        st.sidebar.success(f"COA loaded — {n_accts} accounts, {n_cls} classes")

    if pdf_files:
        for pf in pdf_files:
            txt = extract_pdf_text(pf)
            if txt.strip():
                pdf_texts.append(txt)
        st.sidebar.success(f"{len(pdf_texts)} PDF(s) extracted")

    if bs_balance_file:
        with st.spinner("Parsing Balance Sheet balances..."):
            bs_balances = parse_balance_report(bs_balance_file)
        if bs_balances is not None:
            month_cols = [c for c in bs_balances.columns if c != "Account"]
            st.sidebar.success(f"BS balances loaded — {len(bs_balances)} accounts, {len(month_cols)} months")
        else:
            st.sidebar.warning("Could not parse BS balance file. Check format: accounts as rows, months as columns.")


    # ── Guard ──
    if gl is None:
        st.markdown("""
        <div style="text-align:center; padding:4rem 2rem;">
            <p style="font-size:1.1rem; color:#6E6E73;">
                Upload a <strong>General Ledger</strong> file in the sidebar to get started.
            </p>
            <p style="font-size:0.9rem; color:#8E8E93;">
                Supported formats: CSV, XLSX (QuickBooks Desktop export).
            </p>
        </div>
        """, unsafe_allow_html=True)
        return

    # ── Tabs ──
    tabs = st.tabs([
        "1 · Flux Narrative",
        "2 · Missing Bills",
        "3 · Suspense Reclass",
        "4 · Materiality",
        "5 · IIF Pre-Flight",
        "6 · Reconciliation",
        "7 · BS Top 20 Variance",
        "8 · P&L Top 20 Variance",
    ])

    with tabs[0]:
        report_flux(gl)

    with tabs[1]:
        report_vendor_gap(gl, closing_month, start_month)

    with tabs[2]:
        report_suspense(gl, coa)

    with tabs[3]:
        report_materiality(gl, threshold)

    with tabs[4]:
        checks = report_iif_preflight(gl, coa)

    with tabs[5]:
        report_reconciliation(gl, pdf_texts)

    with tabs[6]:
        report_bs_variance(gl, closing_month, start_month, bs_balances)

    with tabs[7]:
        report_pl_variance(gl, closing_month, start_month)

    # ── Word Document Export ──
    if export_docx:
        st.markdown("---")
        section("Word Document Export")
        if st.button("Generate All Reports as .docx", type="primary"):
            with st.spinner("Generating Apple-branded Word documents..."):
                try:
                    zip_buf, generated = export_all_reports_docx(gl, coa, threshold, pdf_texts, closing_month, start_month, bs_balances)
                    st.success(f"Generated {len(generated)} report(s).")
                    st.download_button(
                        label="Download All Reports (.zip)",
                        data=zip_buf,
                        file_name=f"MonthEnd_Reports_{datetime.date.today().isoformat()}.zip",
                        mime="application/zip",
                    )
                except Exception as e:
                    st.error(f"Export failed: {e}")

    # ── Excel Export ──
    if export_xlsx:
        st.markdown("---")
        section("Excel Workbook Export")
        if st.button("Generate All Reports as .xlsx", type="primary"):
            with st.spinner("Building Excel workbook..."):
                try:
                    xlsx_buf = export_all_reports_xlsx(gl, coa, threshold, pdf_texts, closing_month, start_month, bs_balances)
                    st.success("Excel workbook generated with all 8 reports.")
                    st.download_button(
                        label="Download Excel Workbook (.xlsx)",
                        data=xlsx_buf,
                        file_name=f"MonthEnd_Reports_{datetime.date.today().isoformat()}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    )
                except Exception as e:
                    st.error(f"Excel export failed: {e}")

    # ── IIF Export ──
    if export_iif:
        st.markdown("---")
        section("IIF Export")
        iif_content = generate_iif(gl)
        st.download_button(
            label="Download IIF File",
            data=iif_content,
            file_name=f"month_end_adjustments_{datetime.date.today().isoformat()}.iif",
            mime="text/plain",
        )
        with st.expander("Preview IIF Output"):
            st.code(iif_content[:3000], language="text")

    # Footer
    st.markdown(f"""
    <div class="apple-footer">
        Organization &nbsp;|&nbsp; Finance Department &nbsp;|&nbsp; Month-End Close Agent
        &nbsp;|&nbsp; Generated {datetime.date.today().strftime('%d %b %Y')}
    </div>
    """, unsafe_allow_html=True)


if __name__ == "__main__":
    main()
